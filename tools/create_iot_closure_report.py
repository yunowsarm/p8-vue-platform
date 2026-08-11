from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE


OUT = r"D:\wujie\p8-vue-platform\p8-vue-platform\智慧园区三大模块业务闭环与操作说明.docx"

BLUE = "2E74B5"
NAVY = "1F4D78"
INK = "1F2937"
MUTED = "64748B"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F8FC"
LIGHT_GRAY = "F2F4F7"
GREEN = "13795B"
AMBER = "A16207"
RED = "B42318"
WHITE = "FFFFFF"


def set_run_font(run, size=11, color=INK, bold=False, italic=False):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        edge_data = kwargs.get(edge)
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                element.set(qn("w:{}".format(key)), str(edge_data[key]))


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    grid = tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)


def keep_together(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    p_pr.append(keep)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, 9, MUTED)


def add_styled_paragraph(doc, text="", size=11, color=INK, bold=False, italic=False, before=0, after=6,
                         align=WD_ALIGN_PARAGRAPH.LEFT, style=None):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_label_paragraph(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(label)
    set_run_font(r, 10.5, NAVY, bold=True)
    r = p.add_run(text)
    set_run_font(r, 10.5, INK)
    return p


def add_callout(doc, label, text, fill=PALE_BLUE, label_color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell, 120, 160, 120, 160)
    set_cell_shading(cell, fill)
    border = {"val": "single", "sz": "6", "color": "D8E6F4"}
    set_cell_border(cell, top=border, bottom=border, left=border, right=border)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(label)
    set_run_font(r, 10.5, label_color, bold=True)
    r = p.add_run(text)
    set_run_font(r, 10.5, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_scope_table(doc):
    add_styled_paragraph(doc, "当前页面范围", size=12, color=NAVY, bold=True, before=4, after=4)
    rows = [
        ("智慧安防", "消防设备维保登记", "设备台账、计划、提醒、登记与证据归档"),
        ("智慧安防", "智能周界", "设备与防区、检测规则、告警联动、处置台账"),
        ("智慧安防", "智能巡更", "路线计划、签到核验、异常升级、问题工单"),
        ("能耗管理", "高低压配电能耗", "分级计量、负荷分析、异常用能、定额与报表"),
        ("能耗管理", "照明能耗", "回路策略、异常照明、节能效果与控制审计"),
        ("能耗管理", "给排水能耗", "水平衡、泄漏处置、泵组效率、定额预算"),
        ("智慧通行", "智慧停车", "车位态势、通行/违停、车辆权限、收费与审计"),
        ("智慧通行", "智能门禁", "门区态势、权限任务、异常认证与操作审计"),
        ("智慧通行", "智能访客", "预约审批、临时凭证、联动权限、在园与异常处置"),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [1800, 2400, 5160])
    header = table.rows[0]
    for i, value in enumerate(("模块", "页面", "当前形成的业务链路")):
        cell = header.cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_run_font(r, 10, NAVY, bold=True)
    set_repeat_table_header(header)
    border = {"val": "single", "sz": "4", "color": "D8E1EA"}
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            set_cell_margins(cells[i])
            set_cell_border(cells[i], top=border, bottom=border, left=border, right=border)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, 9.5, INK)
    for cell in header.cells:
        set_cell_border(cell, top=border, bottom=border, left=border, right=border)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_page_block(doc, title, view, operate, closure, evidence, risk=None):
    p = add_styled_paragraph(doc, title, size=13, color=BLUE, bold=True, before=12, after=4, style="Heading 2")
    keep_together(p)
    add_label_paragraph(doc, "可查看信息：", view)
    add_label_paragraph(doc, "怎么操作：", operate)
    add_label_paragraph(doc, "形成闭环：", closure)
    add_label_paragraph(doc, "留痕/输出：", evidence)
    if risk:
        add_label_paragraph(doc, "关键控制：", risk)


def add_section_heading(doc, text):
    p = add_styled_paragraph(doc, text, size=16, color=BLUE, bold=True, before=16, after=6, style="Heading 1")
    keep_together(p)
    return p


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1


def set_page_setup(doc):
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("智慧园区管理平台 · 业务闭环与操作说明")
        set_run_font(r, 9, MUTED)
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("业务汇报材料  |  第 ")
        set_run_font(r, 9, MUTED)
        add_page_field(p)
        r = p.add_run(" 页")
        set_run_font(r, 9, MUTED)


def build_document():
    doc = Document()
    configure_styles(doc)
    set_page_setup(doc)
    section = doc.sections[0]

    # Cover / memo masthead
    add_styled_paragraph(doc, "智慧园区管理平台", size=12, color=BLUE, bold=True, before=40, after=10,
                         align=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, "三大模块业务闭环\n与操作说明", size=25, color=NAVY, bold=True, before=0, after=12,
                         align=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, "智慧安防 · 能耗管理 · 智慧通行", size=13, color=MUTED, before=0, after=30,
                         align=WD_ALIGN_PARAGRAPH.CENTER)
    meta = doc.add_table(rows=3, cols=2)
    set_table_geometry(meta, [2200, 7160])
    meta_rows = [("文档用途", "甲方汇报、演示讲解与验收沟通"), ("覆盖页面", "智慧安防 3 页、能耗管理 3 页、智慧通行 3 页"), ("整理日期", "2026 年 7 月 29 日 · V1.0")]
    for row, (label, value) in zip(meta.rows, meta_rows):
        for cell in row.cells:
            set_cell_margins(cell, 120, 160, 120, 160)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        set_cell_shading(row.cells[1], PALE_BLUE)
        for cell, txt, is_label in ((row.cells[0], label, True), (row.cells[1], value, False)):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(txt)
            set_run_font(r, 10.5, NAVY if is_label else INK, bold=is_label)
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    add_callout(doc, "汇报结论：", "页面已把“数据展示、规则配置、人工处置、联动/派单、复核关闭、记录追溯”串成可演示的业务闭环；当前为前端仿真数据与交互，生产上线时需接入硬件、后端接口及审计服务。", fill="EDF6F3", label_color=GREEN)
    doc.add_page_break()

    add_section_heading(doc, "一、总体闭环与汇报口径")
    add_styled_paragraph(doc, "三大模块采用同一业务框架：设备或终端回传数据后，系统集中展示态势并按规则识别异常；工作人员在详情页完成确认、审批、派单、联动或关闭；最终保留台账、证据、操作记录和统计报表，支持追溯。", size=11, color=INK, after=8)
    add_callout(doc, "统一闭环：", "数据回传 → 态势展示/规则研判 → 告警或待办 → 人工确认、审批或下发 → 联动、派单与现场处置 → 复核关闭 → 台账、审计、报表追溯。")
    add_scope_table(doc)
    add_styled_paragraph(doc, "建议在甲方汇报时先说明“数据来源与责任边界”：硬件、控制器、摄像机、车牌机、表计、移动终端负责采集/执行；平台负责汇聚、规则、流程、协同、记录与分析。", size=10.5, color=MUTED, italic=True, before=6, after=4)

    add_section_heading(doc, "二、智慧安防")
    add_callout(doc, "模块定位：", "从“设备健康与风险发现”延伸至“现场处置、复核关闭和长期留痕”，覆盖消防维保、周界防护与人员巡更三类安防作业。")
    add_page_block(
        doc, "2.1 消防设备维保登记",
        "消防设备总数、在用正常、30 天内到期、已逾期、本月已维保、故障待修；按设备类型的年度维保进度、近期到期清单、设备台账、维保计划与维保记录。",
        "从“设备台账”检索设备并进入登记维保；在“维保计划”查看范围、频次和责任人；在“到期提醒”对到期设备立即登记；填写检查项、维保结果、问题说明、现场图片/视频、费用和下次维保日期。",
        "设备台账 → 计划生成 → 到期提醒 → 现场检查与证据上传 → 结果登记 → 自动形成下次维保节点；故障设备可转维修任务并回写状态。",
        "设备档案、检查项目、执行人/单位、现场凭证、维保结果、下次维保日期和到期状态。",
        "到期未登记持续提醒；维保记录用于后续履历追溯，不能以无记录方式删除。"
    )
    add_page_block(
        doc, "2.2 智能周界",
        "园区地图上的围墙、围栏、出入口、防区，摄像机、雷达、红外对射等设备在线/离线/故障/布撤防状态；实时告警、告警证据、规则列表和设备健康。",
        "在周界总览点击防区、设备或告警进入详情；在“检测规则”配置区域、时段、目标、停留时间、灵敏度和识别来源；在告警详情执行确认、误报标记、派单、复核、关闭以及联动摄像机、声光和移动端。",
        "设备/视频/雷达数据 → 规则识别越界、翻越、入侵、徘徊等行为 → 保存位置、抓拍和前后视频 → 自动联动与人工处置 → 复核关闭 → 告警台账永久留痕。",
        "防区状态、设备健康、告警等级与时间、触发设备、抓拍/视频证据、联动执行、处置人和处置时间线。",
        "对绿植、小动物、雨雪雾、逆光、施工临时撤防可配置区域/时段/灵敏度；告警不支持无记录删除。"
    )
    add_page_block(
        doc, "2.3 智能巡更",
        "当日计划、执行中任务、完成率、异常告警、待处置问题、在线终端；路线地图、当前点、签到记录、异常告警、问题工单。",
        "在“巡更计划”配置路线、点位、班次、人员、时间窗口、检查内容和必拍要求；巡更人员通过二维码、NFC、蓝牙或定位签到并上传图文、语音、视频；发现问题后在问题工单中上报并跟踪。",
        "路线/计划配置 → 移动终端签到与位置校验 → 自动识别漏巡、迟巡、提前巡、越序与异常停留 → 分级通知班组长/值班负责人 → 问题生成安保或维修工单 → 现场处置、复核关闭。",
        "任务进度、人员当前位置、点位签到时间与方式、距离/时效校验、异常类型、现场资料、关联工单和升级记录。",
        "签到校验用于防代扫和事后补签；异常升级规则可在配置页维护。"
    )

    add_section_heading(doc, "三、能耗管理")
    add_callout(doc, "模块定位：", "以分级计量为基础，把“采集、分析、预警、处置、节能优化和报表追溯”连成节能管理闭环。")
    add_page_block(
        doc, "3.1 高低压配电能耗",
        "今日用电、实时功率、最大需量、变压器负载率、本月电费、待处置异常；分级计量下钻、负荷曲线、用电流向、异常用能、定额预算、日/月/季/年报表。",
        "在用电总览按园区、变压器、馈线、回路和表计逐级查看；切换日/月查看负荷曲线与对应横坐标；在异常用能中打开事件并处置；在定额预算维护对标维度，在报表中预览或导出。",
        "表计/配电数据 → 分级模型与负荷分析 → 识别损耗、三相不平衡、功率因数偏低、负荷突变、零值、倒退或离线 → 分析/派单/确认 → 处置归档 → 定额与报表复盘。",
        "原始采集指标、电量流向、趋势曲线、异常说明、关联表计、处置状态、预算执行、同类对象对标及报表。",
        "实际生产需明确电价、峰平谷与财务对账口径，并由后端保存原始采集值和审计轨迹。"
    )
    add_page_block(
        doc, "3.2 照明能耗",
        "今日照明用电、开启回路、平均调光度、策略节电率、异常亮灯、控制器异常；分区趋势、照明策略与环境关联、策略节能对比、控制审计。",
        "在回路表计查看开关、调光和运行时长；在策略节能中配置分区、定时、调光、人感、日历和日出日落条件；在异常照明中处理无人长亮、非使用时段亮灯、能耗突增和状态电量不一致；在控制审计查看人工覆盖。",
        "电量/开关/环境数据 → 策略执行与效果比较 → 发现异常照明 → 人工确认或策略调整 → 对比节电量、节电率和费用 → 控制审计与长期追溯。",
        "分区及回路状态、照度/人员存在、实施前后能耗、节电量与费用变化、控制对象、操作人、原因、优先级和结果。",
        "消防应急照明不纳入普通节能控制；消防及现场硬件控制优先于经审批人工覆盖，人工覆盖优先于自动策略。"
    )
    add_page_block(
        doc, "3.3 给排水能耗",
        "今日用水、瞬时流量、水量平衡差额、泵组耗电、预算执行、疑似泄漏；园区水平衡树图、瞬时流量趋势、泄漏处置、泵组效率和用水定额。",
        "从用水总览查看入口到楼栋/用途的分级水量及未计量差额；在泄漏处置中打开事件、现场核查并派工；在泵组效率中对比单位输水电耗；在定额预算中维护用水目标与预测。",
        "水表/电表/泵控数据 → 水量平衡与夜间最小流量分析 → 识别持续小流量、突增、倒流、离线和长时间补水 → 严重泄漏自动建工单或发起关阀申请 → 人工确认/现场处置 → 复核关闭与预算复盘。",
        "分级水量、流量与阈值、平衡差额、泵组功率/效率、疑似损失、关联工单、供水保护状态和处置结果。",
        "涉及消防、应急或重要生产用水时，平台仅允许现场核查和派单，禁止远程关阀。"
    )

    add_section_heading(doc, "四、智慧通行")
    add_callout(doc, "模块定位：", "围绕“人、车、门、梯、访客凭证”的准入、使用、异常和回收，实现通行数据与业务权限的协同闭环。")
    add_page_block(
        doc, "4.1 智慧停车",
        "车位总数、空余车位、占用率、今日入场、未处置违停、设备异常；分区车位卡片、选中分区占用情况、出入口实时通行、设备健康、违停处置、车辆权限、收费对账与操作审计。",
        "在停车总览选择分区查看占用、预约、违停和周转；点击出入记录查看车辆通行；在违停处置查看车牌与图片，执行语音提醒、通知巡查或转工单；在车辆权限维护白/黑名单、有效期、区域与特殊车位；在操作审计登记人工改牌、抬杆或离线放行。",
        "车位/相机/道闸数据 → 空位引导与通行记录 → 识别违停、堵塞、逆行等异常 → 保存车牌、全景、位置及持续时间 → 语音提醒/派单/处置 → 驶离或关闭 → 收费、对账和审计留痕。",
        "车位占用、区域趋势、车辆类型、进出时间、违停证据、收费/退款状态、设备状态、人工放行与改牌原因。",
        "网络中断、无牌、污损/相似车牌和识别错误应走人工处置；人工操作必须记录人员与原因。"
    )
    add_page_block(
        doc, "4.2 智能门禁",
        "今日通行、在线门点、异常事件、离线/防拆、待同步权限、门长开；门区地图、实时认证与异常、权限管理、下发任务和操作审计。",
        "在门禁总览选择门区/门点查看在线、认证和告警；在权限管理按人员、组织、门区、星期、时段和有效期配置权限；在下发任务检查成功、失败与离线待同步；在事件详情执行确认、远程锁定或处置关闭；在操作审计发起远程开门、常开或批量调整。",
        "人员状态/权限配置 → 控制器任务下发 → 门禁认证与通行记录 → 识别非法认证、权限过期、连续失败、强开、门未关、离线或防拆 → 视频/人工处置 → 审批复核/关闭 → 全程审计。",
        "门区与控制器状态、认证人员/方式、关联视频、权限有效期、下发进度、失败原因、远程操作审批和结果。",
        "远程开门、紧急开门、常开和批量调权需要专门权限、二次确认和审计；重要区域支持审批或双人复核。"
    )
    add_page_block(
        doc, "4.3 智能访客",
        "今日预约、当前在园、待审批、今日离场、异常提醒、外协施工；当前在园访客地图、预约审批、在园管理、异常处置、权限联动。",
        "在预约审批中搜索预约并完成审批、驳回或取消；审批通过后查看二维码/临时凭证、车辆、门区和楼层范围；在在园管理查看实时位置及超时未离场；在异常处置查看未审批到访、黑名单、重复使用、超范围或超时事件并升级或关闭；在权限联动核对停车、门禁和梯控下发状态。",
        "预约登记 → 被访人审批 → 限时/限区/限次凭证 → 停车、门禁、梯控权限下发 → 到访核验与在园管理 → 离场、取消或到期自动回收 → 异常升级和在园清单留存。",
        "访客/团体信息、来访事由、被访人、有效时间、车辆、凭证、审批记录、在园位置、通行范围、联动状态及异常处置时间线。",
        "访客信息应最小化采集、脱敏展示和限权导出；紧急疏散可使用当前在园清单。"
    )

    add_section_heading(doc, "五、统一操作方式与甲方演示路径")
    add_styled_paragraph(doc, "建议每个页面按“先看态势、再看规则、再做处置、最后看留痕”的顺序演示。这样能直观看到不是单纯展示看板，而是可落到人员和工单的业务闭环。", size=11, color=INK, after=6)
    steps = [
        "进入模块总览：先说明顶部指标所代表的待办、风险、设备状态和实时量。",
        "点击业务导航：切换到规则、计划、台账、异常、权限或报表页，说明配置入口。",
        "搜索或筛选记录：按设备、区域、人员、车辆、访客、时间或状态快速定位。",
        "打开详情抽屉/弹窗：展示证据、关联设备、状态变化、时间线和可执行的处理动作。",
        "执行一次典型动作：例如确认告警、派工、审批、登记、重试下发、策略调整或关闭。",
        "返回台账或审计：检查状态、处置人、原因、证据和后续节点已回写，完成闭环说明。",
    ]
    for item in steps:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        r = p.add_run(item)
        set_run_font(r, 10.5, INK)
    add_callout(doc, "演示重点：", "优先选择一条有异常、有待办或有证据的数据完成全流程讲解。例如周界翻越告警、消防设备到期维保、夜间泄漏、违停车辆、访客待审批或门禁非法认证。", fill="FFF7E6", label_color=AMBER)

    add_section_heading(doc, "六、当前版本边界与上线对接清单")
    add_callout(doc, "当前版本说明：", "页面内数据、状态变化、图表、详情、配置和处置均用于前端仿真演示。当前可验证交互和闭环设计，不能替代真实硬件控制、后台持久化、消息推送或生产审计。", fill="FDF1F1", label_color=RED)
    add_styled_paragraph(doc, "建议与甲方确认的上线事项", size=12, color=NAVY, bold=True, before=6, after=4)
    checklist = [
        "硬件清单与数据协议：设备编号、状态字典、实时数据频率、图片/视频取流、控制指令和离线补传规则。",
        "流程与权限：告警等级、派单对象、升级时限、远程控制审批、双人复核范围、误报标记与关闭口径。",
        "工单与消息：安保、维修、客服/值班的工单接口，短信/企业微信/移动端推送，以及回执字段。",
        "财务与报表：停车收费、优惠、退款、发票、用电电价、峰平谷、能耗预算和对账标准。",
        "数据治理：访客个人信息最小化、脱敏、保存期限、导出权限；告警、维保、人工操作不可无记录删除。",
        "验收场景：每类模块至少选取 1 条正常链路、1 条异常链路和 1 条离线/失败恢复链路进行联调验收。",
    ]
    for item in checklist:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        r = p.add_run(item)
        set_run_font(r, 10.5, INK)

    doc.core_properties.title = "智慧园区三大模块业务闭环与操作说明"
    doc.core_properties.subject = "智慧安防、能耗管理、智慧通行业务闭环汇报材料"
    doc.core_properties.author = "智慧园区管理平台"
    doc.core_properties.comments = "当前前端仿真版业务闭环与操作说明"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
